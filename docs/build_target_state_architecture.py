from __future__ import annotations

import os
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "agentic-pdlc-target-state-reference-architecture.docx"
ASSET_DIR = Path(tempfile.mkdtemp(prefix="agentic-pdlc-target-state-"))

NAVY = "0B2545"
BLUE = "2563EB"
TEAL = "0F766E"
VIOLET = "7C3AED"
AMBER = "D97706"
GREEN = "15803D"
RED = "B91C1C"
SLATE = "475569"
MUTED = "64748B"
LIGHT = "F1F5F9"
PALE_BLUE = "EFF6FF"
PALE_TEAL = "ECFDF5"
PALE_AMBER = "FFF7ED"
WHITE = "FFFFFF"
INK = "0F172A"
BORDER = "CBD5E1"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_table_borders(table, color=BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, side="left", color=BLUE, size=18, space=8):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text_node, end])
    set_run_font(run, size=8.5, color=MUTED)
    return run


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.10

    specs = {
        "Title": (30, NAVY, 0, 8),
        "Subtitle": (14, SLATE, 0, 14),
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    # Remove decorative residue inherited from Word's built-in Title style.
    title_ppr = styles["Title"]._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.50)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.167

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(8.5)
    caption.font.color.rgb = rgb(MUTED)
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True


def configure_section(section, first=False):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = first


def configure_landscape_section(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.30)
    section.footer_distance = Inches(0.30)
    section.different_first_page_header_footer = False


def set_running_furniture(section, content_width=6.5):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = ""
    p.style = "Normal"
    p_pr = p._p.get_or_add_pPr()
    old_tabs = p_pr.find(qn("w:tabs"))
    if old_tabs is not None:
        p_pr.remove(old_tabs)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(content_width), WD_TAB_ALIGNMENT.RIGHT)
    left = p.add_run("AGENTIC PDLC CONTEXT FRAMEWORK")
    set_run_font(left, size=8, color=MUTED, bold=True)
    p.add_run("\t")
    right = p.add_run("TARGET-STATE REFERENCE ARCHITECTURE")
    set_run_font(right, size=8, color=MUTED)
    set_paragraph_border(p, side="bottom", color=BORDER, size=5, space=4)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.style = "Normal"
    fp_pr = fp._p.get_or_add_pPr()
    old_footer_tabs = fp_pr.find(qn("w:tabs"))
    if old_footer_tabs is not None:
        fp_pr.remove(old_footer_tabs)
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.paragraph_format.space_before = Pt(3)
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(content_width), WD_TAB_ALIGNMENT.RIGHT)
    lr = fp.add_run("Leadership Review | Future-State Vision")
    set_run_font(lr, size=8.2, color=MUTED)
    fp.add_run("\t")
    pr = fp.add_run("Page ")
    set_run_font(pr, size=8.2, color=MUTED)
    add_field(fp, "PAGE")


def add_body(doc, text, *, bold_lead=None, italic=False, after=6, keep=False, size=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_together = keep
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.right_indent = Inches(0)
    p.paragraph_format.first_line_indent = Inches(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=size, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=size, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size, italic=italic)
    return p


def add_bullet(doc, text, *, level=0, bold_lead=None):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(5)
    if bold_lead and text.startswith(bold_lead):
        run = p.add_run(bold_lead)
        set_run_font(run, bold=True)
        run2 = p.add_run(text[len(bold_lead):])
        set_run_font(run2)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def new_numbering_instance(doc):
    numbering = doc.part.numbering_part.element
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), "7")
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def add_numbered(doc, text, *, bold_lead=None, num_id=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    if num_id is not None:
        p_pr = p._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num])
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_numbered_group(doc, items):
    num_id = new_numbering_instance(doc)
    for item in items:
        add_numbered(doc, item, num_id=num_id)


def add_callout(doc, label, text, *, fill=PALE_BLUE, accent=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.10)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, side="left", color=accent, size=22, space=8)
    r1 = p.add_run(label.upper() + "  ")
    set_run_font(r1, size=9.2, color=accent, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.2, color=INK, bold=False)
    return p


def add_kicker(doc, text, color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.first_line_indent = Inches(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text.upper())
    set_run_font(r, size=9.5, color=color, bold=True)
    return p


def add_table(doc, headers, rows, widths_dxa, *, header_fill=NAVY, first_col_bold=False, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    table.rows[0].height = None
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(header)
        set_run_font(r, size=9.0, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            if row_idx % 2 == 1:
                set_cell_shading(cells[col_idx], "F8FAFC")
            p = cells[col_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK, bold=(first_col_bold and col_idx == 0))
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)
    return table


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    set_paragraph_shading(p, LIGHT)
    set_paragraph_border(p, side="left", color=SLATE, size=12, space=6)
    r = p.add_run(text)
    set_run_font(r, name="Consolas", size=8.4, color=INK)
    return p


def add_page_break(doc):
    if doc.sections[-1].orientation == WD_ORIENT.LANDSCAPE:
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        configure_section(section, first=False)
        set_running_furniture(section)
    else:
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)


def add_figure(doc, number, title, image_path, caption, takeaway, *, intro=None):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_landscape_section(section)
    set_running_furniture(section, content_width=9.70)
    add_kicker(doc, f"Architecture view {number}", TEAL)
    doc.add_heading(title, level=1)
    if intro:
        add_body(doc, intro, after=7, size=9.8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_together = True
    picture = p.add_run().add_picture(str(image_path), width=Inches(9.62))
    picture._inline.docPr.set(
        "descr",
        f"Architecture diagram: {title}. {caption}",
    )
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(f"Figure {number}. {caption}")
    add_callout(doc, "Leadership takeaway", takeaway, fill=PALE_TEAL, accent=TEAL)


def render_dot(name, source):
    dot_path = ASSET_DIR / f"{name}.dot"
    png_path = ASSET_DIR / f"{name}.png"
    dot_path.write_text(source, encoding="utf-8")
    subprocess.run(
        ["dot", "-Tpng", "-Gdpi=220", str(dot_path), "-o", str(png_path)],
        check=True,
        env={**os.environ, "LANG": "en_US.UTF-8"},
    )
    return png_path


DOT_HEADER = f'''
digraph G {{
  graph [bgcolor="white", fontname="Arial", fontsize=16, pad=0.22, nodesep=0.28, ranksep=0.42, splines=ortho];
  node [shape=box, style="rounded,filled", fontname="Arial", fontsize=11.5, margin="0.14,0.10", color="#{BORDER}", penwidth=1.3, fillcolor="white", fontcolor="#{INK}"];
  edge [fontname="Arial", fontsize=9.5, color="#{SLATE}", fontcolor="#{SLATE}", penwidth=1.3, arrowsize=0.72];
'''


def build_diagrams():
    diagrams = {}

    diagrams[1] = render_dot("01_value_stream", DOT_HEADER + f'''
      rankdir=TB;
      gov [shape=box, style="rounded,filled", fillcolor="#{NAVY}", color="#{NAVY}", fontcolor="white", label="GOVERNANCE & CONTROL PLANE\nWorkflow orchestration | State & checkpoints | Policy | Human authority | Budget | Plugin registry"];
      s1 [fillcolor="#{PALE_BLUE}", color="#{BLUE}", label="1  DEMAND & REQUIREMENTS\nIntent / incident / portfolio trigger\nRequirement + acceptance baseline\nAmbiguity and policy analysis"];
      s2 [fillcolor="#{PALE_BLUE}", color="#{BLUE}", label="2  ARCHITECTURE & DESIGN\nDesign agent + architecture retrieval\nExplained impact + threat model\nADR + test strategy"];
      s3 [fillcolor="#F5F3FF", color="#{VIOLET}", label="3  CHANGE REALIZATION\nImplementation agent pool\nEphemeral worktree + build analysis\nActual-change containment"];
      s4 [fillcolor="#{PALE_TEAL}", color="#{TEAL}", label="4  QUALITY & ASSURANCE\nRisk-based test obligations\nTest data broker + isolated runners\nEvidence collection"];
      s5 [fillcolor="#{PALE_TEAL}", color="#{TEAL}", label="5  RELEASE & OPERATE\nProgressive delivery + attestation\nObservability + rollback\nIncident and defect intake"];
      fabric [shape=box, style="rounded,filled", fillcolor="#{TEAL}", color="#{TEAL}", fontcolor="white", label="CONTEXT & EVIDENCE FABRIC\nTemporal context graph | Hybrid retrieval | Evidence/artifact store | Immutable audit ledger"];
      integrations [shape=box, style="rounded,filled,dashed", fillcolor="#F8FAFC", color="#{SLATE}", label="ENTERPRISE INTEGRATION PLANE\nRequirements | SCM | CI/CD | Test management | CMDB | Observability | ITSM"];
      {{rank=same; s1; s2; s3; s4; s5}}
      gov -> s3 [label="governed work order", color="#{NAVY}"];
      s1 -> s2 [label="G1  approved criteria"];
      s2 -> s3 [label="G2  authorized impact boundary"];
      s3 -> s4 [label="actual revision pair"];
      s4 -> s5 [label="G3  evidence + signed decision"];
      s3 -> fabric [label="queries / assertions", dir=both, color="#{TEAL}"];
      s5 -> fabric [label="outcomes", style=dashed, color="#{TEAL}"];
      fabric -> integrations [label="references + events", dir=both, style=dashed];
      s5 -> s1 [label="incident / drift / opportunity", style=dashed, color="#{RED}", constraint=false];
    }}''')

    diagrams[2] = render_dot("02_context_fabric", DOT_HEADER + f'''
      rankdir=TB;
      sources [fillcolor="#F8FAFC", color="#{SLATE}", label="AUTHORITATIVE ENTERPRISE SOURCES\nBusiness: requirements | portfolio | controls\nEngineering: Git | APIs | build graph | CMDB\nAssurance: tests | CI | telemetry | incidents"];
      acquire [fillcolor="#F5F3FF", color="#{VIOLET}", label="ACQUIRE & NORMALIZE\nBYO connectors: webhook | poll | reconcile\nAnalyzers: AST/LSP | contracts | runtime | documents\nClassification | redaction | policy filters"];
      trust [fillcolor="#{NAVY}", color="#{NAVY}", fontcolor="white", label="IDENTITY, ONTOLOGY & TRUST\nResolve tenant | project | system | revision\nValidate typed semantics and schema versions\nStamp provenance | confidence | freshness | dedupe"];
      fabric [shape=cylinder, fillcolor="#{PALE_TEAL}", color="#{TEAL}", label="POLYGLOT CONTEXT & EVIDENCE FABRIC\nImmutable assertion ledger\nTemporal topology and domain graph\nEvidence / artifact object store\nHybrid retrieval and evaluation stores"];
      services [fillcolor="#{PALE_BLUE}", color="#{BLUE}", label="CONTEXT & DECISION SERVICES\nAs-of query | traceability | hybrid grounding\nCanonical impact and test obligations\nReadiness | policy context | evidence explorer"];
      consumers [fillcolor="#F8FAFC", color="#{SLATE}", label="GOVERNED CONSUMERS\nLifecycle agents | Developers | Approvers | Auditors\nAPIs and event subscribers\nEvaluation and analytics"];
      {{rank=same; sources; acquire; trust}}
      {{rank=same; fabric; services; consumers}}
      sources -> acquire [label="events | revisions | observations"];
      acquire -> trust [label="candidate assertions"];
      trust -> fabric [label="append validated assertion", color="#{TEAL}"];
      fabric -> services [label="revision-consistent projections"];
      services -> consumers [label="answer + provenance + gaps"];
      consumers -> fabric [label="evidence + outcomes", style=dashed, color="#{TEAL}", constraint=false];
      trust -> services [label="identity + semantics", style=dashed, color="#{NAVY}"];
    }}''')

    diagrams[3] = render_dot("03_temporal_model", DOT_HEADER + f'''
      rankdir=TB;
      event [shape=note, fillcolor="#F8FAFC", label="OBSERVED DOMAIN EVENT\nRequirement changed | Code indexed\nTest executed | Release deployed"];
      envelope [fillcolor="#{PALE_BLUE}", color="#{BLUE}", label="ASSERTION ENVELOPE\nTenant | project | system | stable entity\nPredicate | object | source identity\nRevision / run / attempt | schema version\nValid time | observed time | confidence | evidence ref"];
      identity [fillcolor="#F5F3FF", color="#{VIOLET}", label="IDENTITY RESOLUTION\nStable logical entity\nImmutable entity revision\nImmutable execution attempt"];
      ledger [shape=cylinder, fillcolor="#{PALE_TEAL}", color="#{TEAL}", label="BITEMPORAL ASSERTION LEDGER\nAppend | supersede | retract\nNever rewrite historical evidence"];
      snapshot [fillcolor="#{NAVY}", color="#{NAVY}", fontcolor="white", label="VALIDATED TOPOLOGY SNAPSHOT\nProject + revision set | provider versions\nQuality vector | checksum | activation state"];
      views [fillcolor="#{PALE_TEAL}", color="#{TEAL}", label="MATERIALIZED, PURPOSE-SPECIFIC VIEWS\nCurrent validated graph | as-of-commit graph\nAs-of-run trace | readiness | evaluation history"];
      consumers [fillcolor="#{PALE_BLUE}", color="#{BLUE}", label="REPRODUCIBLE CONSUMERS\nAgent grounding | Impact | QA | Audit\nEvidence explorer | Analytics"];
      valid [fillcolor="#{PALE_AMBER}", color="#{AMBER}", label="VALID TIME\nWhen true in source"];
      system [fillcolor="#{PALE_AMBER}", color="#{AMBER}", label="SYSTEM TIME\nWhen observed and stored"];
      selector [fillcolor="#{PALE_AMBER}", color="#{AMBER}", label="QUERY SELECTOR\nProject + revision + as-of time"];
      {{rank=same; event; envelope; identity}}
      {{rank=same; ledger; snapshot; views; consumers}}
      {{rank=same; valid; system; selector}}
      event -> envelope [label="normalize + stamp"];
      envelope -> identity [label="resolve"];
      identity -> ledger [label="append assertion"];
      ledger -> snapshot [label="validated fact set"];
      snapshot -> views [label="atomic activation"];
      ledger -> views [label="evidence history", style=dashed];
      views -> consumers [label="answer + complete lineage"];
      valid -> system -> selector [color="#{AMBER}"];
      selector -> views [label="temporal scope", color="#{AMBER}"];
    }}''')

    diagrams[4] = render_dot("04_impact_engine", DOT_HEADER + f'''
      rankdir=TB;
      change [fillcolor="#F8FAFC", color="#{SLATE}", label="1  SEMANTIC CHANGESET\nImmutable base + head revisions\nAdd | modify | rename | delete\nSymbol | API | schema | config | infrastructure"];
      semantic [fillcolor="#F5F3FF", color="#{VIOLET}", label="2  SEED RESOLUTION\nLanguage-aware parsing and compatibility\nMap changed artifacts to entities and owners\nSurface every unmapped change"];
      signals [fillcolor="#{PALE_TEAL}", color="#{TEAL}", label="3  MULTI-SIGNAL PROPAGATION\nStatic calls/imports | API/event/data contracts\nRuntime traces | build/deployment topology\nRequirements/ownership/history | test/defect evidence\nProvider-defined direction, trigger and budget semantics"];
      fusion [fillcolor="#{NAVY}", color="#{NAVY}", fontcolor="white", label="4  SIGNAL FUSION & PATH SCORING\nChange severity | business criticality | path strength\nEvidence class | provenance | freshness | confidence\nTraversal budget | client risk policy | historical outcomes"];
      tiers [fillcolor="#{PALE_AMBER}", color="#{AMBER}", label="5  EXPLAINED IMPACT TIERS\nT0 Directly changed | T1 Must inspect / must test\nT2 Probable downstream | T3 Notify / monitor\nUNKNOWN Unmapped, stale or insufficient evidence"];
      consumers [fillcolor="#{PALE_BLUE}", color="#{BLUE}", label="6  ONE CANONICAL IMPACT ASSESSMENT\nDesign authorization | implementation containment\nRegression and test-data scope | release-risk decision\nReason paths | obligations | blind spots | residual risk"];
      eval [fillcolor="#F5F3FF", color="#{VIOLET}", label="7  OUTCOME CALIBRATION\nFailures | escapes | incidents | overrides\nRuntime behavior | false blocks | missed tests\nMeasured, versioned and governed adjustment"];
      {{rank=same; change; semantic; signals; fusion}}
      {{rank=same; tiers; consumers; eval}}
      change -> semantic [label="actual revision pair"];
      semantic -> signals [label="typed seeds"];
      signals -> fusion [label="paths + evidence"];
      fusion -> tiers [label="risk + unknowns"];
      tiers -> consumers [label="immutable assessment"];
      consumers -> eval [label="observed outcomes", style=dashed, color="#{VIOLET}"];
      eval -> fusion [label="approved calibration", style=dashed, color="#{VIOLET}", constraint=false];
    }}''')

    diagrams[5] = render_dot("05_quality_data", DOT_HEADER + f'''
      rankdir=TB;
      impact [fillcolor="#{NAVY}", color="#{NAVY}", fontcolor="white", label="1  IMPACT ASSESSMENT\nAffected entities + explanation paths\nRisk tiers + unknowns\nMandatory assurance obligations"];
      planner [fillcolor="#F5F3FF", color="#{VIOLET}", label="2  ASSURANCE PLANNER\nResolve existing assets before generating\nReuse | author gaps | exploratory charters\nPrioritize by tier, criticality and policy"];
      data_plan [fillcolor="#{PALE_AMBER}", color="#{AMBER}", label="3  SCENARIO-SCOPED DATA PLAN\nEntities | privacy class | volume\nIsolation | lifetime | fidelity\nSetup and teardown obligations"];
      broker [fillcolor="#{PALE_TEAL}", color="#{TEAL}", label="4  GOVERNED TEST DATA BROKER\nSynthetic factory | masked/minimized subset\nTokenized clone | service virtualization\nIdempotent DataLease | TTL | cleanup attestation"];
      catalog [fillcolor="#{PALE_BLUE}", color="#{BLUE}", label="VERSIONED TEST & COVERAGE CATALOG\nScenario | script | framework | owner\nDeclared + static + runtime coverage\nReliability | defects | execution history"];
      runner [fillcolor="#F5F3FF", color="#{VIOLET}", label="5  ISOLATED EXECUTION PLANE\nEphemeral environment and worktree\nParallel runners | fault injection\nRevision-bound attempt + lease identity"];
      evidence [fillcolor="#{PALE_TEAL}", color="#{TEAL}", label="6  EVIDENCE COLLECTOR\nTest leaves | coverage | traces | artifacts\nData lease + teardown proof | defects\nSigned, attributable evidence bundle"];
      gate [shape=diamond, fillcolor="#{PALE_AMBER}", color="#{AMBER}", label="7  DETERMINISTIC\nQA DECISION"];
      ledger [shape=cylinder, fillcolor="#{TEAL}", color="#{TEAL}", fontcolor="white", label="EVIDENCE LEDGER\nAttempt + observation\nDecision + attestation"];
      {{rank=same; impact; planner; data_plan; broker}}
      {{rank=same; catalog; runner; evidence; gate; ledger}}
      impact -> planner [label="obligations"];
      catalog -> planner [label="eligible assets + coverage", style=dashed];
      planner -> data_plan [label="data requirements"];
      data_plan -> broker [label="authorized strategy"];
      planner -> runner [label="execution plan"];
      broker -> runner [label="scoped DataLease"];
      runner -> evidence [label="observed results"];
      evidence -> gate [label="signed bundle"];
      evidence -> ledger [label="append", style=dashed];
      gate -> ledger [label="policy decision"];
      ledger -> catalog [label="coverage + reliability history", style=dashed, constraint=false];
    }}''')

    diagrams[6] = render_dot("06_plugin_architecture", DOT_HEADER + f'''
      rankdir=TB;
      lifecycle [fillcolor="#{PALE_AMBER}", color="#{AMBER}", label="PLUGIN LIFECYCLE\nDiscover manifest → Validate → Conformance suite → Capability negotiation → Policy approval → Activate → Observe → Roll back"];
      clients [shape=record, style="rounded,filled,dashed", fillcolor="#F8FAFC", label="{{CLIENT ENTERPRISE LANDSCAPE|Requirements & portfolio|SCM & architecture repositories|CI/CD & test platforms|Data, deployment, observability & ITSM}}"];
      inbound [shape=record, style="rounded,filled,dashed", fillcolor="#F5F3FF", color="#{VIOLET}", label="{{INBOUND PORTS|RequirementSource|SourceControl|CodeIntelligence|ContextProvider|IdentityResolver|EventConnector}}"];
      kernel [shape=record, fillcolor="#{NAVY}", color="#{NAVY}", fontcolor="white", label="{{STABLE FRAMEWORK KERNEL|Identity + scope + ontology|Workflow state machine|Canonical Impact Engine|Policy + deterministic gates|Context/evidence contracts|Tenancy + audit}}"];
      outbound [shape=record, style="rounded,filled,dashed", fillcolor="#F5F3FF", color="#{VIOLET}", label="{{OUTBOUND PORTS|LLM / Agent Provider|Test Author / Executor|TestDataProvider|WorkDispatch|BuildDeploy|Evidence / Audit Sink}}"];
      targets [shape=record, style="rounded,filled,dashed", fillcolor="#F8FAFC", label="{{CLIENT / PROVIDER TARGETS|Client-hosted agents|Commercial or local models|GitHub Actions / Jenkins / Kubernetes|Testing and release platforms}}"];
      contract [fillcolor="#{PALE_TEAL}", color="#{TEAL}", label="NON-BYPASSABLE CONTRACT\nVersioned schemas | Idempotency | Correlation | Async completion | Error taxonomy | Health | Telemetry | Classification | Secret references"];
      {{rank=same; clients; inbound; kernel; outbound; targets}}
      lifecycle -> kernel [label="approved capability manifest", color="#{AMBER}"];
      clients -> inbound [label="native API / event"];
      inbound -> kernel [label="canonical command / assertion"];
      kernel -> outbound [label="typed work order"];
      outbound -> targets [label="provider-native execution"];
      targets -> outbound [label="signed result + evidence ref", style=dashed];
      outbound -> kernel [label="normalized outcome", style=dashed];
      contract -> kernel [label="enforced invariants", color="#{TEAL}"];
      contract -> inbound [style=dashed, color="#{TEAL}"];
      contract -> outbound [style=dashed, color="#{TEAL}"];
    }}''')

    diagrams[7] = render_dot("07_security_deployment", DOT_HEADER + f'''
      rankdir=TB;
      users [fillcolor="#F8FAFC", color="#{SLATE}", label="ENTERPRISE USERS & IDENTITY\nAdministrators | Developers | Approvers | Auditors\nEnterprise IdP | MFA | RBAC + attribute policy"];
      control [style="rounded,filled,dashed", penwidth=2, fillcolor="#{NAVY}", color="#{RED}", fontcolor="white", label="CONTROL-PLANE TRUST BOUNDARY\nManaged or client-hosted\nAPI gateway + tenant router\nWorkflow orchestrator + checkpoints\nPolicy engine + human authority\nContext query + provider registry\nAudit, budgets and observability"];
      execution [style="rounded,filled,dashed", penwidth=2, fillcolor="#{PALE_TEAL}", color="#{RED}", label="CLIENT EXECUTION TRUST BOUNDARY\nWorkload identity + signed-job verifier\nEphemeral runner + isolated worktree\nAgent sandbox + test runner\nTest-data broker + evidence collector\nClient-local cache / artifacts"];
      systems [fillcolor="#F8FAFC", color="#{SLATE}", label="CLIENT SYSTEMS & RUNTIME\nGit | requirements | vault | databases\nCI/CD | cloud | Kubernetes\nObservability | CMDB | SIEM"];
      external [style="rounded,filled,dashed", fillcolor="#F5F3FF", color="#{VIOLET}", label="OPTIONAL EXTERNAL PROVIDERS\nCommercial or local model endpoint\nCoding / test / specialist agent\nOnly policy-filtered context crosses boundary"];
      modes [fillcolor="#{PALE_AMBER}", color="#{AMBER}", label="PLACEMENT POLICY\nManaged | Hybrid (enterprise default)\nPrivate | Regulated isolated enclave"];
      security [fillcolor="#{PALE_BLUE}", color="#{BLUE}", label="ZERO-TRUST CONTROLS\nShort-lived credentials | private links | egress allowlist\nSigned work orders and artifacts | secret references\nClassification | minimization | residency | DLP"];
      {{rank=same; users; control; execution; systems}}
      {{rank=same; modes; security; external}}
      users -> control [label="OIDC session + scoped authority"];
      control -> execution [label="signed least-privilege work order", color="#{RED}"];
      execution -> systems [label="short-lived identity + private connectivity"];
      systems -> execution [label="scenario data + runtime evidence", style=dashed, color="#{TEAL}"];
      execution -> control [label="signed metadata + evidence references", style=dashed, color="#{TEAL}"];
      execution -> external [label="approved egress only", style=dashed, color="#{VIOLET}"];
      modes -> execution [label="placement decision", color="#{AMBER}"];
      security -> control [label="enforced policy"];
      security -> execution [label="enforced policy"];
      control -> systems [label="audit events to SIEM", style=dashed, constraint=false];
    }}''')

    diagrams[8] = render_dot("08_learning_loop", DOT_HEADER + f'''
      rankdir=TB;
      observe [fillcolor="#{PALE_TEAL}", color="#{TEAL}", label="1  OPERATE & OBSERVE\nQA outcomes | Telemetry | Escapes\nIncidents | Rollbacks | Overrides\nCost and latency"];
      attribute [fillcolor="#{PALE_BLUE}", color="#{BLUE}", label="2  DIAGNOSE & ATTRIBUTE\nLink outcome to requirement, change,\nimpact path, tests, data, provider,\npolicy and release"];
      eval [fillcolor="#F5F3FF", color="#{VIOLET}", label="3  EVALUATION FACTORY\nHistorical replay | Mutation tests\nImpact and test-selection recall\nSafety, evidence and cost benchmarks"];
      candidate [fillcolor="#F5F3FF", color="#{VIOLET}", label="4  CANDIDATE IMPROVEMENTS\nTests | Coverage mappings | Extractors\nImpact rules | Prompts | Plugins\nPolicy threshold proposals"];
      govern [shape=diamond, fillcolor="#{PALE_AMBER}", color="#{AMBER}", label="5  GOVERN\nHuman review\nChampion / challenger\nSecurity + rollback plan"];
      promote [fillcolor="#{PALE_TEAL}", color="#{GREEN}", label="6  CONTROLLED PROMOTION\nSigned version | Feature flag\nCanary cohort | Staged rollout\nAutomatic rollback"];
      runtime [fillcolor="#{NAVY}", color="#{NAVY}", fontcolor="white", label="7  VERSIONED RUNTIME\nTest library | Graph rules | Policies\nPrompts | Plugins | Orchestration"];
      score [fillcolor="#F8FAFC", color="#{SLATE}", label="MEASURED SCORECARD\nImpact recall at bounded radius\nRegression-selection miss rate\nEscapes / false blocks | evidence completeness\nManual intervention | cost | latency | drift"];
      ledger [shape=cylinder, fillcolor="#{TEAL}", color="#{TEAL}", fontcolor="white", label="Context & evidence fabric\nDataset + decision + version + result"];
      {{rank=same; observe; attribute; eval; candidate}}
      {{rank=same; govern; promote; runtime}}
      {{rank=same; score; ledger}}
      observe -> attribute [label="evidence-linked signals"];
      attribute -> eval [label="labeled cases"];
      eval -> candidate [label="measured failure modes"];
      candidate -> govern [label="proposed version"];
      govern -> promote [label="approved + signed", color="#{GREEN}"];
      promote -> runtime [label="canary / feature flag", color="#{GREEN}"];
      runtime -> observe [label="post-promotion outcomes", style=dashed, color="#{VIOLET}", constraint=false];
      runtime -> promote [label="threshold breach / rollback", color="#{RED}", constraint=false];
      eval -> score [label="publish metrics"];
      attribute -> ledger [style=dashed, color="#{TEAL}"];
      govern -> ledger [style=dashed, color="#{TEAL}"];
      promote -> ledger [style=dashed, color="#{TEAL}"];
    }}''')

    return diagrams


def main():
    diagrams = build_diagrams()
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    configure_section(section, first=True)
    set_running_furniture(section)

    # Cover page
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(62)
    add_kicker(doc, "Future-state reference architecture", TEAL)
    title = doc.add_paragraph(style="Title")
    title.add_run("Agentic PDLC\nContext Framework")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Target-State Vision, Architectural Principles and Enterprise Design")
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(18)
    set_paragraph_border(rule, side="bottom", color=TEAL, size=18, space=4)
    add_callout(
        doc,
        "Executive proposition",
        "A governed context fabric that converts heterogeneous enterprise facts into reproducible impact decisions, bounded agentic work and verifiable delivery evidence.",
        fill=PALE_TEAL,
        accent=TEAL,
    )
    add_body(doc, "Prepared for: Executive Leadership, Enterprise Architecture, Product and Engineering", bold_lead="Prepared for:")
    add_body(doc, "Document status: Target State - Leadership Review", bold_lead="Document status:")
    add_body(doc, "Scope: Vendor-neutral, implementation-agnostic framework vision", bold_lead="Scope:")
    add_body(doc, "Version: 1.0 | 25 August 2026", bold_lead="Version:")
    footer_note = doc.add_paragraph()
    footer_note.paragraph_format.space_before = Pt(42)
    footer_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = footer_note.add_run("CONTEXT BEFORE AUTONOMY  |  EVIDENCE BEFORE RELEASE  |  EXPLAINABILITY BY DESIGN")
    set_run_font(r, size=8.8, color=TEAL, bold=True)

    # Executive summary
    add_page_break(doc)
    add_kicker(doc, "Executive perspective")
    doc.add_heading("The strategic intent", level=1)
    add_callout(
        doc,
        "Leadership message",
        "This platform is not a chain of autonomous agents. It is a governed delivery decision system in which agents propose and execute work, while deterministic controls authorize every consequential state transition.",
    )
    add_body(
        doc,
        "The target framework creates a durable context and evidence layer across the product delivery lifecycle. It connects business intent, architecture, code, tests, release decisions and operational outcomes without replacing the enterprise systems that own those records. The result is a shared, temporal model of why a change exists, what it can affect, what must be tested, what evidence was observed and who authorized progression.",
    )
    add_body(
        doc,
        "The differentiating capability is Impact Intelligence: one canonical, explainable assessment derived from a revision pair and multiple relationship signals. Design, implementation containment, QA selection and release risk all consume that same assessment under different policy lenses. Uncertainty and missing context remain visible rather than being converted into false confidence.",
    )
    doc.add_heading("What leadership is enabling", level=2)
    for text, lead in [
        ("Faster flow with bounded autonomy. Agents work inside explicit scope, budget, tool and data boundaries.", "Faster flow with bounded autonomy."),
        ("Risk-based assurance. Test obligations and release controls follow the actual change and its explained downstream consequences.", "Risk-based assurance."),
        ("Enterprise adaptability without client forks. Connectors, analyzers, agents, test frameworks and deployment tools are replaceable behind governed contracts.", "Enterprise adaptability without client forks."),
        ("Audit-ready decisions. Every important outcome is reproducible from the project, revision, snapshot, policy, execution attempt and evidence set.", "Audit-ready decisions."),
        ("A learning system with controls. Improvements are evaluated, versioned, approved, canaried and reversible; agents do not silently rewrite their own guardrails.", "A learning system with controls."),
    ]:
        add_bullet(doc, text, bold_lead=lead)

    doc.add_heading("Leadership decisions requested", level=2)
    add_numbered_group(doc, [
        "Endorse the context fabric and canonical Impact Engine as strategic platform capabilities, not utilities embedded inside individual agents.",
        "Fund framework correctness - identity, temporal truth, provider contracts, evaluation and security - ahead of broad connector or agent breadth.",
        "Adopt a hybrid execution posture in which sensitive code, credentials and test data can remain inside the client boundary.",
        "Establish design partners and measurable proof thresholds for impact recall, test-selection misses, evidence completeness and cycle-time improvement.",
    ])

    # Document map
    add_page_break(doc)
    add_kicker(doc, "Document map")
    doc.add_heading("How to read this architecture", level=1)
    add_body(doc, "The document moves from strategic intent to the core semantic model, then zooms into impact intelligence, assurance execution, pluggability, security and the governed learning loop.")
    sections = [
        ("1", "Vision and outcomes", "Why the framework exists and the capabilities it must create."),
        ("2", "Architectural tenets", "The non-negotiable rules that protect the future platform."),
        ("3", "End-to-end reference architecture", "How governed work flows from intent through operations."),
        ("4", "Context and evidence fabric", "How enterprise truth becomes temporal, trusted context."),
        ("5", "Semantic and temporal model", "Identity, revisions, assertions, snapshots and evidence."),
        ("6", "Canonical Impact Engine", "How one explainable assessment drives every downstream decision."),
        ("7", "Quality, test data and evidence", "How test obligations become isolated execution and auditable proof."),
        ("8", "BYO plugin framework", "How clients bring tools and implementations without platform forks."),
        ("9", "Security and deployment", "How execution follows data across SaaS, hybrid and regulated environments."),
        ("10", "Evaluation, governance and roadmap", "How the platform improves safely and scales as a product."),
    ]
    add_table(doc, ["Section", "Topic", "Purpose"], sections, [900, 2500, 5960], first_col_bold=True)
    add_callout(doc, "Reading principle", "The diagrams describe stable framework responsibilities and contracts. Products, vendors, storage technologies and models remain replaceable.", fill=PALE_AMBER, accent=AMBER)

    # Vision and outcomes
    add_page_break(doc)
    add_kicker(doc, "1 | Vision and outcomes")
    doc.add_heading("North-star vision", level=1)
    add_callout(doc, "Vision", "A governed, temporally correct and extensible context fabric that turns enterprise facts into reproducible impact decisions and verifiable product-delivery evidence.", fill=PALE_TEAL, accent=TEAL)
    add_body(doc, "The framework provides the decision substrate for an end-to-end agentic PDLC. It does not attempt to replace Jira, Git, CI/CD, test management, observability, CMDB or ITSM. It resolves their identities, preserves their revisions and relationships, and exposes trusted context through common decision contracts.")
    doc.add_heading("Target business outcomes", level=2)
    outcomes = [
        ("Delivery flow", "Reduce handoff latency and repetitive analysis", "Bounded agent work, reusable context, deterministic progression"),
        ("Change safety", "Find consequential downstream effects early", "Multi-signal impact paths, uncertainty and risk tiers"),
        ("Quality economics", "Run the smallest defensible test portfolio", "Change-linked test obligations and measured selection recall"),
        ("Auditability", "Explain why every material decision was made", "Immutable snapshots, evidence and policy decisions"),
        ("Client adaptability", "Integrate heterogeneous enterprise estates", "Capability-based plugins and policy packs"),
        ("Continuous learning", "Improve with observed outcomes", "Versioned evaluation, controlled promotion and rollback"),
    ]
    add_table(doc, ["Outcome", "Enterprise objective", "Framework response"], outcomes, [1700, 3000, 4660], first_col_bold=True)
    doc.add_heading("In scope", level=2)
    for item in [
        "Governed orchestration of requirement, design, implementation, QA, release and operational feedback.",
        "A shared semantic model spanning business intent, system topology, assurance assets and evidence.",
        "Explainable impact, test obligations, release readiness and agent grounding.",
        "A provider SDK for client-owned and third-party integrations, models, agents and execution services.",
        "Security, policy, tenancy, audit, evaluation and lifecycle governance as core platform capabilities.",
    ]:
        add_bullet(doc, item)
    add_page_break(doc)
    add_kicker(doc, "Section 1 / Vision and outcomes - continued")
    doc.add_heading("Explicit non-goals", level=2)
    for item in [
        "Becoming the authoritative system for requirements, source code, incidents or production configuration.",
        "Mandating a specific graph database, model provider, CI/CD platform or test framework.",
        "Allowing probabilistic agents to approve their own high-risk actions or mutate policy without governance.",
        "Claiming universal impact certainty; the platform exposes unsupported surfaces and residual risk.",
    ]:
        add_bullet(doc, item)
    doc.add_heading("Durable kernel, replaceable edge", level=2)
    boundary_rows = [
        ("Framework kernel", "Identity, temporal truth, workflow state, impact contracts, policy decisions, evidence and audit"),
        ("Replaceable providers", "Connectors, analyzers, retrieval engines, models, agents, test frameworks, runners and deployment tools"),
        ("Enterprise authority", "Requirements, source code, configurations, tests, releases, incidents and operational telemetry"),
    ]
    add_table(doc, ["Boundary", "Responsibility"], boundary_rows, [2100, 7260], first_col_bold=True)
    add_callout(doc, "Design boundary", "A client can replace any provider without changing the meaning of an entity, snapshot, assessment, evidence item or decision.", fill=PALE_BLUE, accent=BLUE)

    # Tenets
    add_page_break(doc)
    add_kicker(doc, "2 | Architectural tenets")
    doc.add_heading("The platform constitution", level=1)
    add_body(doc, "These tenets are technology-independent. They should be enforced through domain contracts, conformance tests and architecture decision records.")
    tenets = [
        ("Context before autonomy", "An agent may act only within an explicit project, revision, policy and tool context."),
        ("Deterministic authority", "Agents propose or execute; policy and human authority decide consequential state transitions."),
        ("Immutable historical truth", "Entity revisions, execution attempts, assertions, evidence and decisions are append-only and addressable."),
        ("One impact truth", "Design, implementation, QA and release consume one canonical Impact Assessment."),
        ("Explainability by construction", "Every impact, test obligation and gate outcome carries paths, provenance, policy and known gaps."),
        ("Uncertainty is data", "Unknown, unmapped, stale and low-confidence context remains visible and policy-addressable."),
        ("Behavioral pluggability", "Providers contribute capabilities and semantics, not merely records or labels."),
        ("No client forks", "Client variation is expressed through plugins, policy packs, ontology packages and deployment configuration."),
        ("Source systems remain authoritative", "The platform references and reconciles enterprise truth rather than silently replacing it."),
        ("Execution follows the data", "Sensitive code, credentials and test data can remain within a client-controlled execution boundary."),
        ("Evidence over assertion", "Release confidence is based on observed, revision-bound evidence rather than generated narrative."),
        ("Learning is governed", "Improvements pass evaluation, approval, versioning, canary and rollback controls."),
    ]
    add_table(doc, ["Tenet", "Architectural implication"], tenets, [2500, 6860], first_col_bold=True, font_size=8.8)

    add_figure(
        doc, 1, "Governed end-to-end delivery value stream", diagrams[1],
        "Agents perform lifecycle work while deterministic gates, shared context and evidence govern transitions.",
        "The platform is a governed delivery operating system. Every transition is tied to a project, revision, policy decision and evidence set.",
        intro="The reference flow preserves familiar lifecycle stages but places a common control plane above them and a temporal context fabric beneath them.",
    )

    # Logical architecture
    add_page_break(doc)
    add_kicker(doc, "3 | End-to-end reference architecture")
    doc.add_heading("Logical capability model", level=1)
    add_body(doc, "The framework is organized into stable platform layers. Client tools and providers integrate through explicit boundaries; the kernel retains identity, policy, evidence and decision semantics.")
    layers = [
        ("Experience and actors", "Administrator, developer, approver, auditor and product-owner journeys; APIs and event interfaces", "Core experience + client extensions"),
        ("Workflow and governance", "Orchestration, checkpoints, budgets, human authority, policy evaluation and resumability", "Framework core"),
        ("Decision intelligence", "Impact analysis, assurance planning, release readiness, explanations and residual-risk decisions", "Framework core + policy plugins"),
        ("Context services", "Identity, ontology, snapshots, traceability, hybrid retrieval, evidence and audit", "Framework core + storage adapters"),
        ("Agent services", "Requirement, design, implementation, test and operations agents", "Replaceable providers"),
        ("Execution plane", "Ephemeral workspaces, build, test, test data, deployment and evidence collection", "Client-hosted or managed"),
        ("Integration plane", "Requirements, SCM, CI/CD, catalogs, CMDB, observability and ITSM", "BYO connectors"),
        ("Platform foundation", "Tenancy, security, secrets, telemetry, cost controls, availability and data lifecycle", "Framework core + infrastructure adapters"),
    ]
    add_table(doc, ["Layer", "Primary responsibilities", "Ownership model"], layers, [1900, 5200, 2260], first_col_bold=True, font_size=8.8)
    doc.add_heading("Core control flow", level=2)
    for text, lead in [
        ("Intent is baselined. Requirements and acceptance criteria become versioned, resolvable context.", "Intent is baselined."),
        ("Design is authorized. Proposed scope is evaluated against explained impact, architecture policy and test strategy.", "Design is authorized."),
        ("Actual change is re-evaluated. The produced revision pair, not the design narrative, becomes the source for containment and QA.", "Actual change is re-evaluated."),
        ("Evidence is observed. Tests, coverage, security analysis and deployment signals attach to immutable execution attempts.", "Evidence is observed."),
        ("Release is decided. Policy consumes the same context, assessment and evidence bundle and records a signed decision.", "Release is decided."),
        ("Operations close the loop. Incidents, rollbacks and runtime behavior become labeled outcomes for evaluation.", "Operations close the loop."),
    ]:
        add_bullet(doc, text, bold_lead=lead)

    # Context fabric
    add_figure(
        doc, 2, "Enterprise context and evidence fabric", diagrams[2],
        "Authoritative enterprise records are normalized into immutable assertions and materialized into purpose-specific context projections.",
        "Agents never become systems of record. The fabric resolves identity, preserves evidence and returns answers with provenance, confidence and known gaps.",
        intro="The fabric is deliberately polyglot: temporal relationships, artifacts, retrieval indexes and evaluation data have different access patterns but share one identity and assertion contract.",
    )

    add_page_break(doc)
    add_kicker(doc, "4 | Context and evidence fabric")
    doc.add_heading("Core responsibilities", level=1)
    responsibilities = [
        ("Scope and identity", "Resolve tenant, project, system, logical entity and immutable revision identities."),
        ("Ontology and semantics", "Validate entity/relationship types, extension packages, propagation semantics and schema compatibility."),
        ("Assertion ingestion", "Authorize producers; validate, classify, stamp provenance, deduplicate and append facts."),
        ("Snapshot lifecycle", "Build, validate and atomically activate complete topology projections."),
        ("Evidence lifecycle", "Append execution observations, artifacts, defects, overrides and policy decisions without rewriting history."),
        ("Context query", "Serve current and as-of views, explanation paths, traceability and quality qualifiers."),
        ("Hybrid retrieval", "Combine semantic, lexical, graph and metadata filters while remaining revision-bound."),
        ("Reconciliation", "Detect source drift and represent correction through supersede/retract assertions."),
        ("Evaluation", "Create labeled datasets and measure provider, retrieval, impact and test-selection performance."),
    ]
    add_table(doc, ["Capability", "Responsibility"], responsibilities, [2200, 7160], first_col_bold=True)
    doc.add_heading("Truth classes", level=2)
    truth_rows = [
        ("Authoritative", "Git commit, requirement revision, deployment, test result", "Source-system identity and deep link", "Imported; source remains authoritative"),
        ("Derived", "Static import, API relationship, ownership inference", "Provider version, snapshot and confidence", "Immutable within a topology snapshot"),
        ("Observed", "Runtime call, trace coverage, incident, rollback", "Execution or observation identity", "Append-only evidence"),
        ("Inferred", "LLM-proposed relationship, predicted severity", "Model/prompt/tool versions and confidence", "Reviewable; never silently authoritative"),
    ]
    add_table(doc, ["Truth class", "Examples", "Minimum provenance", "Lifecycle"], truth_rows, [1400, 2750, 2550, 2660], first_col_bold=True, font_size=8.4)
    add_callout(doc, "Non-negotiable", "A probabilistic inference, compiler-derived relationship and runtime observation must never appear indistinguishable to a consumer or policy.", fill=PALE_AMBER, accent=AMBER)

    # Temporal model
    add_figure(
        doc, 3, "Temporal identity, assertion and snapshot model", diagrams[3],
        "Stable entities are separated from immutable revisions and execution attempts; bitemporal assertions produce reproducible views.",
        "The framework must answer both what is true now and what was known for a specific revision when a decision was made.",
        intro="Historical truth is preserved by separating logical identity, revision identity and assertion identity. Updates append superseding facts rather than mutating old evidence.",
    )

    add_page_break(doc)
    add_kicker(doc, "5 | Semantic and temporal model")
    doc.add_heading("Canonical domain objects", level=1)
    model_rows = [
        ("ScopeContext", "Mandatory tenant, project, system, repository, revisions, snapshot, policy and execution scope."),
        ("Entity", "Stable logical identity for a requirement, service, API, artifact, test, environment or release."),
        ("EntityRevision", "Immutable version of an entity at a source revision or content hash."),
        ("AssertionEnvelope", "Typed subject-predicate-object statement with provenance, time, confidence and evidence reference."),
        ("TopologySnapshot", "Validated, complete set of derived topology assertions for an explicit project and revision set."),
        ("ChangeSet", "Typed comparison between base and head snapshots, including additions, modifications, moves and deletions."),
        ("ImpactAssessment", "Immutable, explainable consequence model produced from ChangeSet, snapshot and policy versions."),
        ("ExecutionAttempt", "Immutable performance of a workflow, build, test or deployment at a specific revision."),
        ("EvidenceObservation", "Observed result or artifact bound to an execution attempt and relevant entity revisions."),
        ("PolicyDecision", "Signed, reproducible authorization or denial with rule version, inputs and authority."),
    ]
    add_table(doc, ["Object", "Purpose"], model_rows, [2200, 7160], first_col_bold=True, font_size=8.8)
    doc.add_heading("Mandatory scope contract", level=2)
    add_code_block(doc, "ScopeContext\n  tenant_id            required\n  project_id           required\n  system_ids[]         explicit\n  repository_ids[]     explicit\n  base_revision        immutable\n  head_revision        immutable\n  topology_snapshot_id immutable\n  policy_version       immutable\n  execution_attempt_id optional until execution")
    doc.add_heading("Snapshot activation invariant", level=2)
    add_numbered_group(doc, [
        "Build a complete candidate snapshot and all required projections in an isolated staging scope.",
        "Validate schema, provenance, identity integrity, provider health, coverage and quality thresholds.",
        "Sign the snapshot manifest and compute content checksums.",
        "Atomically move the active pointer; a failed build leaves the previous snapshot untouched.",
        "Retain historical snapshots according to policy and preserve every assessment's referenced snapshot.",
    ])

    # Impact engine
    add_figure(
        doc, 4, "Multi-signal canonical Impact Engine", diagrams[4],
        "A revision pair is semantically analyzed, expanded across typed relationship signals and converted into explained impact tiers.",
        "There is no magical blast-radius number. The platform produces tiered consequences, reason paths and explicit unknowns that every lifecycle consumer shares.",
        intro="The engine fuses static, contract, runtime, deployment, business and quality signals. Client policy controls how evidence classes propagate and when uncertainty blocks progression.",
    )

    add_page_break(doc)
    add_kicker(doc, "6 | Canonical Impact Engine")
    doc.add_heading("Impact as a first-class decision artifact", level=1)
    add_body(doc, "The Impact Engine is the platform's primary differentiator. It converts a concrete revision pair into a reproducible assessment. Different lifecycle stages may apply different policies to the assessment, but they must not maintain independent traversal algorithms or hidden dependency models.")
    doc.add_heading("Processing stages", level=2)
    stages = [
        ("1. Semantic change analysis", "Identify changed symbols, contracts, schemas, configuration, infrastructure and behavior; preserve adds, deletes and renames."),
        ("2. Seed resolution", "Resolve changed artifacts to modules, services, APIs, requirements and ownership; isolate unmapped changes."),
        ("3. Typed propagation", "Traverse relationships using provider-defined direction, change triggers, budgets and evidence class."),
        ("4. Signal fusion", "Combine path strength, business criticality, source quality, freshness, runtime evidence and historical outcomes."),
        ("5. Tiering and obligations", "Assign direct, mandatory, probable, monitoring and unknown tiers; synthesize test and review obligations."),
        ("6. Policy evaluation", "Apply client risk tolerance, regulated controls, criticality and human-approval thresholds."),
        ("7. Explanation", "Return complete reason paths, provenance, blind spots, algorithm version and residual risk."),
    ]
    add_table(doc, ["Stage", "Design responsibility"], stages, [2500, 6860], first_col_bold=True, font_size=8.8)
    doc.add_heading("Impact tiers", level=2)
    tiers = [
        ("T0 - Direct", "Changed entity or contract", "Always inspect; direct tests and owners"),
        ("T1 - Mandatory", "High-confidence consequence or policy-critical dependency", "Must inspect and must satisfy test obligation"),
        ("T2 - Probable", "Material downstream likelihood with adequate evidence", "Risk-based test or explicit acceptance"),
        ("T3 - Monitor", "Weak, long-range or operational relationship", "Notify, observe or include under elevated risk"),
        ("Unknown", "Unmapped surface, stale context or insufficient provider coverage", "Fail closed, widen scope or require human acceptance"),
    ]
    add_table(doc, ["Tier", "Meaning", "Default treatment"], tiers, [1700, 3850, 3810], first_col_bold=True)
    doc.add_heading("Minimum Impact Assessment contract", level=2)
    add_code_block(doc, "ImpactAssessment\n  assessment_id, scope, change_set_id\n  snapshot_id, algorithm_version, policy_version\n  changed_entities[]\n  affected_entities[] { tier, risk, reasons[] }\n  explanation_paths[] { nodes, edges, provider, evidence_class, freshness }\n  test_obligations[] { target, rationale, coverage_requirement }\n  unknowns[] { surface, reason, prescribed_action }\n  quality_vector, residual_risk, generated_at")
    add_callout(doc, "Policy rule", "No empty result is silently treated as no impact. Every changed artifact must be mapped, declared irrelevant by policy or returned as an explicit unknown.", fill=PALE_AMBER, accent=AMBER)

    # QA and data
    add_figure(
        doc, 5, "Risk-based quality, test data and evidence architecture", diagrams[5],
        "Impact-derived obligations drive test planning, scenario-scoped data leases, isolated execution and an immutable evidence bundle.",
        "Test selection, test data and release evidence are one assurance chain. Setup and teardown are governed artifacts, not hidden test-script side effects.",
        intro="Quality execution starts with the same Impact Assessment used by design and implementation. The assurance planner resolves existing tests, generates only genuine gaps and obtains data through a governed broker.",
    )

    add_page_break(doc)
    add_kicker(doc, "7 | Quality, test data and evidence")
    doc.add_heading("From impact to test obligation", level=1)
    for text, lead in [
        ("Select before generating. Resolve versioned tests with observed coverage and reliability before authoring new assets.", "Select before generating."),
        ("Coverage is revision-bound. Distinguish declared coverage, static linkage and runtime-observed coverage.", "Coverage is revision-bound."),
        ("Every impacted entity has a disposition. Tested, inspected, monitored, accepted or explicitly unknown.", "Every impacted entity has a disposition."),
        ("Evidence is attached at the smallest defensible level. Scenario, script, test leaf, impacted entity, environment and execution attempt.", "Evidence is attached at the smallest defensible level."),
        ("A successful transport is not a passing gate. The deterministic QA decision is an explicit signed artifact.", "A successful transport is not a passing gate."),
    ]:
        add_bullet(doc, text, bold_lead=lead)
    doc.add_heading("Test data as a governed lease", level=2)
    data_rows = [
        ("Synthetic generation", "Default for common functional paths and privacy-sensitive domains", "Schema-valid, scenario-specific, reproducible seed"),
        ("Masked subset", "Complex production-like relationships where synthetic fidelity is insufficient", "Approved extraction, tokenization, minimization and expiry"),
        ("Virtualized dependency", "Unavailable, costly or destructive downstream systems", "Contract-versioned behavior and recorded simulation profile"),
        ("Ephemeral environment", "Stateful integration and destructive workflow testing", "Isolated namespace, infrastructure lease and teardown proof"),
        ("Reference fixture", "Stable deterministic unit/component cases", "Versioned content hash and ownership"),
    ]
    add_table(doc, ["Strategy", "Use when", "Required controls"], data_rows, [1900, 3600, 3860], first_col_bold=True, font_size=8.6)
    doc.add_heading("Setup and teardown protocol", level=2)
    protocol = [
        "The planner emits a DataRequirement with entities, privacy class, volume, isolation and lifetime.",
        "The broker selects an authorized provider and returns an idempotent DataLease bound to the execution attempt.",
        "The runner verifies lease readiness and records the data version before executing tests.",
        "Teardown is invoked regardless of test outcome; TTL and compensating cleanup cover runner failure.",
        "The broker emits teardown attestation. Missing cleanup proof is a policy-visible assurance gap.",
    ]
    add_numbered_group(doc, protocol)

    # Plugin architecture
    add_figure(
        doc, 6, "Behavioral BYO plugin architecture", diagrams[6],
        "Client systems and execution targets integrate through versioned inbound and outbound ports around a non-bypassable framework kernel.",
        "Clients bring implementations, not forks. Identity, state, evidence, policy and audit semantics remain stable regardless of provider choice.",
        intro="Pluggability is governed at activation time. Providers declare capabilities, permissions, versions, quality and operational behavior and must pass the framework conformance suite.",
    )

    add_page_break(doc)
    add_kicker(doc, "8 | BYO and ecosystem architecture")
    doc.add_heading("Provider categories", level=1)
    provider_rows = [
        ("Context acquisition", "Requirements, SCM, architecture catalogs, CMDB, documents, events", "Canonical assertions and source references"),
        ("Relationship intelligence", "Language analyzers, build graphs, API/schema/event parsers, runtime topology", "Typed relationships plus propagation semantics"),
        ("Agent capability", "Requirement, design, coding, test, security and operations agents", "Typed work order and signed outcome"),
        ("Assurance execution", "Test frameworks, environments, data providers, scanners", "Execution evidence and cleanup attestation"),
        ("Delivery execution", "Work dispatch, build, artifact, deployment and rollback platforms", "Attested artifact and deployment observation"),
        ("Storage and retrieval", "Graph, ledger, object, search, vector and analytics implementations", "Versioned storage/query contracts"),
        ("Policy and governance", "Client policy packs, control catalogs, approval and audit sinks", "Deterministic decision extensions"),
    ]
    add_table(doc, ["Provider family", "Examples", "Framework output"], provider_rows, [1900, 4200, 3260], first_col_bold=True, font_size=8.6)
    doc.add_heading("Capability manifest", level=2)
    add_code_block(doc, "ProviderManifest\n  provider_id, version, contract_versions[]\n  capabilities[] { entities, relationships, commands }\n  supported_stacks[]\n  relationship_semantics[] { direction, triggers, confidence, obligations }\n  data_classes[], permissions[], secret_references[]\n  execution_model, locality, health_contract\n  incremental_support, idempotency, error_taxonomy\n  telemetry_contract, quality_metrics, rollback_version")
    doc.add_heading("Provider activation lifecycle", level=2)
    add_numbered_group(doc, [
        "Discover and validate the signed provider manifest.",
        "Resolve configuration and secrets by reference; never persist raw credentials in plugin configuration.",
        "Run contract, security, tenancy, replay, idempotency and failure-mode conformance tests.",
        "Negotiate capabilities and compare them with project policy and required coverage.",
        "Approve, activate and observe the provider with an explicit rollback target.",
        "Block or degrade safely when health, quality, cost or drift thresholds are breached.",
    ])
    add_callout(doc, "Extension rule", "Client-defined ontology types become useful only when their provider supplies validation, traversal, explanation and policy semantics.", fill=PALE_AMBER, accent=AMBER)

    # Security
    add_figure(
        doc, 7, "Hybrid deployment and zero-trust execution", diagrams[7],
        "A governed control plane dispatches signed work into a client-controlled execution plane using short-lived identity and policy-filtered context.",
        "Execution follows the data. Source code, credentials and sensitive test data can remain inside the client's governed boundary.",
        intro="The architecture supports managed, hybrid and isolated deployment without changing the framework contracts. Placement is a project policy, not an application fork.",
    )

    add_page_break(doc)
    add_kicker(doc, "9 | Security, tenancy and deployment")
    doc.add_heading("Security model", level=1)
    controls = [
        ("Identity", "Enterprise OIDC, workload identity, MFA, short-lived credentials and strong execution-attempt identity"),
        ("Authorization", "Tenant/project scope, RBAC plus attribute policy, separation of duties and tool-level permission boundaries"),
        ("Data protection", "Classification, minimization, redaction, encryption, residency, retention and evidence-reference indirection"),
        ("Execution isolation", "Ephemeral workspaces, sandboxing, egress allowlists, signed work orders and attested artifacts"),
        ("Secrets", "Vault references, per-job retrieval, rotation and prohibition on persistence in prompts, logs or graph assertions"),
        ("Supply chain", "Signed plugins, SBOM, provenance, conformance, vulnerability policy and controlled rollback"),
        ("Audit", "Immutable command, tool, model, evidence, approval and policy-decision records forwarded to client SIEM"),
        ("AI safety", "Tool allowlists, context filtering, prompt-injection defenses, output validation, budget and human-authority boundaries"),
    ]
    add_table(doc, ["Control domain", "Target-state design"], controls, [1900, 7460], first_col_bold=True, font_size=8.8)
    doc.add_heading("Deployment profiles", level=2)
    profiles = [
        ("Managed", "Managed control and execution planes", "Low sensitivity; fastest onboarding", "Logical tenant isolation plus customer-specific policy"),
        ("Hybrid", "Managed control plane; client-hosted execution and data", "Enterprise default", "Code, credentials and test data remain client-side"),
        ("Private", "Fully deployed in client cloud or data center", "Regulated or sovereign workloads", "Client controls infrastructure, keys and operations"),
        ("Isolated enclave", "Disconnected or tightly allowlisted deployment", "Highest restriction", "Local models/providers; curated update bundles"),
    ]
    add_table(doc, ["Profile", "Placement", "Typical fit", "Boundary"], profiles, [1300, 2600, 2500, 2960], first_col_bold=True, font_size=8.4)

    # Learning loop
    add_figure(
        doc, 8, "Governed closed-loop learning", diagrams[8],
        "Observed delivery outcomes become labeled evaluation cases; only measured, approved and signed improvements reach runtime.",
        "The framework learns through controlled promotion, never through agents silently rewriting their own controls.",
        intro="The evaluation factory is a product subsystem. It measures whether impact, test selection, providers, prompts and policies improve without increasing risk or creating tenant leakage.",
    )

    add_page_break(doc)
    add_kicker(doc, "10 | Evaluation and governance")
    doc.add_heading("Accuracy is an operating discipline", level=1)
    add_body(doc, "Impact and test-selection credibility cannot be established by architecture alone. Each client and technology estate requires calibrated evidence. The framework therefore treats evaluation datasets, replay and quality thresholds as production capabilities.")
    score_rows = [
        ("Impact recall at bounded radius", "How much observed consequence is captured within an acceptable review/test budget", "Time-split history, mutation, failures and incidents"),
        ("Test-selection miss rate", "Required tests omitted from the selected portfolio", "Full-suite shadow runs and production escapes"),
        ("False-block rate", "Changes delayed without corresponding risk reduction", "Gate overrides and post-release outcomes"),
        ("Evidence completeness", "Obligations with revision-bound, attributable and signed evidence", "Evidence graph conformance"),
        ("Context quality vector", "Coverage, freshness, unresolved surfaces, provider health and hash equality", "Snapshot manifest and provider telemetry"),
        ("Agent containment", "Tool, scope, data and budget violations or prevented attempts", "Control-plane audit"),
        ("Flow and economics", "Cycle time, manual intervention, compute/model cost and test runtime", "Workflow and execution telemetry"),
    ]
    add_table(doc, ["Measure", "Question answered", "Primary evidence"], score_rows, [2200, 4300, 2860], first_col_bold=True, font_size=8.4)
    doc.add_heading("Governance operating model", level=2)
    governance = [
        ("Platform architecture council", "Core contracts, ADRs, compatibility and strategic guardrails"),
        ("Ontology and identity stewardship", "Canonical semantics, extension-package review and entity resolution"),
        ("Provider owners", "Capability quality, conformance, security posture, support and lifecycle"),
        ("Policy owners", "Risk thresholds, regulated controls, human authority and exception handling"),
        ("Evaluation and model risk", "Datasets, metrics, bias/leakage checks, champion/challenger decisions"),
        ("Platform SRE and security", "Availability, isolation, incident response, supply chain and observability"),
        ("Client administrators", "Project configuration, provider activation, policy selection and access"),
    ]
    add_table(doc, ["Role", "Accountability"], governance, [2600, 6760], first_col_bold=True, font_size=8.8)

    # NFRs and decisions
    add_page_break(doc)
    add_kicker(doc, "11 | Quality attributes and architecture decisions")
    doc.add_heading("Target quality attributes", level=1)
    nfrs = [
        ("Reproducibility", "Every assessment and decision can be replayed from immutable scope, snapshot, algorithm and policy versions."),
        ("Tenant isolation", "All identities, storage, queries, jobs, telemetry and learned artifacts are tenant/project scoped by construction."),
        ("Availability", "Query and decision services degrade predictably; asynchronous snapshot builds do not corrupt active context."),
        ("Scalability", "Incremental providers, partitioned snapshots, bounded traversals, materialized projections and asynchronous long-running analysis."),
        ("Interoperability", "Versioned APIs/events, provider SDKs, standard evidence envelopes and backward-compatible capability negotiation."),
        ("Observability", "End-to-end correlation across workflow, provider, model, tool, snapshot, evidence and policy decision."),
        ("Portability", "Storage, model, agent, runner and deployment implementations remain replaceable."),
        ("Resilience", "Idempotent work orders, checkpointed orchestration, retry policies, compensating cleanup and signed replay."),
        ("Cost governance", "Per-project budgets, model/tool quotas, caching, test portfolio optimization and showback."),
    ]
    add_table(doc, ["Attribute", "Design response"], nfrs, [1900, 7460], first_col_bold=True, font_size=8.8)
    doc.add_heading("Foundational architecture decisions", level=2)
    adrs = [
        ("ADR-001", "Context fabric over monolithic graph", "Use one semantic contract with fit-for-purpose ledger, graph, retrieval, object and analytical projections."),
        ("ADR-002", "Bitemporal immutable assertions", "Preserve historical truth and enable as-of-source and as-of-observation queries."),
        ("ADR-003", "Build-and-activate snapshots", "Prevent partial state and make every consumer revision-consistent."),
        ("ADR-004", "Canonical Impact Assessment", "Eliminate divergent blast-radius logic across lifecycle stages."),
        ("ADR-005", "Behavioral provider capabilities", "Support client extension without core changes or semantic ambiguity."),
        ("ADR-006", "Policy as a versioned decision artifact", "Separate agent reasoning from organizational authority."),
        ("ADR-007", "Hybrid execution plane", "Keep sensitive client assets in the client trust boundary."),
        ("ADR-008", "Governed learning", "Require evaluation and approval before runtime changes."),
    ]
    add_table(doc, ["Decision", "Position", "Why it matters"], adrs, [1300, 3100, 4960], first_col_bold=True, font_size=8.4)

    # Roadmap
    add_page_break(doc)
    add_kicker(doc, "12 | Adoption roadmap")
    doc.add_heading("Build the spine before the breadth", level=1)
    add_body(doc, "The roadmap sequences irreversible framework contracts before expanding provider breadth or agent sophistication. Each horizon has an outcome-based exit criterion.")
    roadmap = [
        ("H0 - Contract foundation", "Identity, ScopeContext, ontology registry, assertion envelope, snapshots, evidence and provider manifest", "Contracts versioned; conformance harness proves isolation, replay and compatibility"),
        ("H1 - Context spine", "Source connectors, assertion ledger, topology/domain projections, query APIs, hybrid retrieval and audit", "One revision-consistent snapshot supports explainable trace and grounding"),
        ("H2 - Impact and assurance", "Semantic ChangeSet, multi-signal Impact Engine, test obligations, test data broker and evidence gate", "Design, implementation, QA and release use one assessment; unknowns fail safely"),
        ("H3 - Enterprise scale", "Hybrid deployment, policy packs, regulated controls, additional analyzers, runtime topology and operations links", "Design partners meet isolation, quality and operational SLOs"),
        ("H4 - Ecosystem and learning", "Provider marketplace, evaluation factory, champion/challenger calibration and reusable domain packs", "Partners add capability without forks; improvements are measured and governed"),
    ]
    add_table(doc, ["Horizon", "Primary scope", "Exit criterion"], roadmap, [1650, 4300, 3410], first_col_bold=True, font_size=8.5)
    doc.add_heading("Proof strategy", level=2)
    add_numbered_group(doc, [
        "Select design partners with different stacks, compliance profiles and delivery toolchains.",
        "Establish a time-split historical corpus and full-suite shadow execution before claiming selection accuracy.",
        "Replay known defects and incidents to prove explanation paths and severity-sensitive recall.",
        "Demonstrate a provider replacement without changing core workflow or decision semantics.",
        "Demonstrate as-of audit: reconstruct a release decision using the exact requirement, revision, snapshot, evidence and policy versions.",
        "Demonstrate client-bound execution with no sensitive source, credentials or test data persisted in the managed control plane.",
    ])
    doc.add_heading("Investment guardrails", level=2)
    add_callout(doc, "Priority", "Invest first in temporal correctness, one impact contract, conformance, evidence and evaluation. Connector count and UI breadth do not compensate for a weak decision substrate.", fill=PALE_AMBER, accent=AMBER)

    # Risks
    add_page_break(doc)
    add_kicker(doc, "13 | Strategic risks and mitigations")
    doc.add_heading("Risks to manage deliberately", level=1)
    risks = [
        ("False confidence", "Incomplete context appears authoritative", "Expose evidence class, freshness, coverage and unknowns; fail safely by policy"),
        ("Ontology sprawl", "Client extensions fragment semantics", "Governed packages, namespace ownership, compatibility and semantic conformance"),
        ("Client forks", "Custom integrations alter core behavior", "Capability contracts, policy packs and strict kernel boundaries"),
        ("Accuracy theater", "Proxy metrics become marketing claims", "Client-specific datasets, confidence intervals, shadow runs and outcome-linked measures"),
        ("Data leakage", "Code or test data crosses an unintended boundary", "Hybrid execution, minimization, egress policy, DLP and reference-based evidence"),
        ("Agent overreach", "Probabilistic output changes governed state", "Deterministic state machine, tool authorization, human authority and immutable audit"),
        ("Provider lock-in", "Models or tools become embedded in semantics", "Provider-neutral contracts, capability negotiation and replay tests"),
        ("Feedback bias", "Learning reinforces incomplete historical patterns", "Curated datasets, mutation/escape cases, challenger evaluation and tenant isolation"),
        ("Cost and latency", "Analysis breadth erodes delivery economics", "Budgets, bounded traversal, caching, async analysis and value-based evidence policies"),
    ]
    add_table(doc, ["Risk", "Failure mode", "Architectural mitigation"], risks, [1700, 3000, 4660], first_col_bold=True, font_size=8.5)

    # Leadership close
    add_page_break(doc)
    add_kicker(doc, "14 | Leadership alignment")
    doc.add_heading("The decisions that shape the platform", level=1)
    add_callout(doc, "Recommendation", "Treat the Context Fabric, Impact Engine, Policy Engine, Evidence Ledger and Provider SDK as the product's durable kernel. Treat every agent, model, connector, analyzer, test framework and deployment tool as replaceable.", fill=PALE_TEAL, accent=TEAL)
    decisions = [
        ("Strategic boundary", "Approve the framework as a governed decision platform rather than an agent-automation application."),
        ("Product investment", "Prioritize identity, temporal truth, impact, evidence, conformance and evaluation ahead of feature breadth."),
        ("Deployment posture", "Adopt hybrid as the enterprise default, with managed and isolated profiles supported by the same contracts."),
        ("Proof standard", "Require outcome-linked accuracy and quality measures before external claims of blast-radius or test-selection effectiveness."),
        ("Operating model", "Create platform, ontology, provider, policy, evaluation and security ownership with explicit decision rights."),
        ("Market strategy", "Recruit design partners whose diversity proves pluggability and whose historical data can calibrate impact intelligence."),
    ]
    add_table(doc, ["Decision area", "Leadership position"], decisions, [2200, 7160], first_col_bold=True)
    doc.add_heading("What success looks like", level=2)
    for item in [
        "A developer can explain every selected test and every omitted test from the same revision-bound Impact Assessment.",
        "An auditor can reconstruct a release decision exactly as it was made, including evidence and policy versions.",
        "A client can replace a source-control, model, test or deployment provider without changing framework semantics.",
        "A new technology stack increases capability through a provider, while unsupported surfaces remain explicit.",
        "Operational outcomes improve impact and assurance models through measured, approved and reversible promotion.",
    ]:
        add_bullet(doc, item)
    close = doc.add_paragraph()
    close.paragraph_format.space_before = Pt(18)
    close.paragraph_format.space_after = Pt(0)
    close.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(close, NAVY)
    r = close.add_run("THE NORTH STAR\nContext before autonomy. Evidence before release. Explainability before scale.")
    set_run_font(r, size=13, color=WHITE, bold=True)

    # Appendix
    add_page_break(doc)
    add_kicker(doc, "Appendix A")
    doc.add_heading("Core framework invariants", level=1)
    invariants = [
        "Every operation is explicitly scoped to tenant, project and relevant immutable revisions.",
        "No historical assertion, execution result, artifact reference or policy decision is overwritten.",
        "Every active topology is a validated snapshot switched atomically with its projections.",
        "Every changed artifact is mapped, explicitly excluded or represented as an unknown.",
        "Every impact and obligation has at least one explanation path or an explicit policy rationale.",
        "Every provider is versioned, capability-declared, permission-scoped, observable and rollback-capable.",
        "Every consequential state transition is authorized by deterministic policy and, where required, named human authority.",
        "Every evidence item is attributable to a revision, execution attempt, environment and producer.",
        "Every learned change is evaluated, approved, versioned, canaried and reversible.",
        "Every client-specific learned artifact remains tenant-isolated unless explicitly governed for sharing.",
    ]
    add_numbered_group(doc, invariants)
    doc.add_heading("Minimum public service contracts", level=2)
    services = [
        ("Assertion Service", "Append, supersede and retract typed assertions; validate provenance and authorization."),
        ("Snapshot Service", "Build, validate, activate, resolve and compare topology snapshots."),
        ("Context Query Service", "Entity, path, trace, as-of and hybrid retrieval queries with quality qualifiers."),
        ("Impact Service", "Create, retrieve, compare and explain immutable Impact Assessments."),
        ("Evidence Service", "Register execution attempts, evidence, artifacts, defects, overrides and attestations."),
        ("Policy Decision Service", "Evaluate and record versioned decisions with complete inputs and authority."),
        ("Provider Registry", "Register capabilities, conformance, configuration, health, activation and rollback."),
        ("Evaluation Service", "Manage datasets, replay, metrics, champion/challenger results and promotion evidence."),
    ]
    add_table(doc, ["Service", "Contract responsibility"], services, [2200, 7160], first_col_bold=True, font_size=8.8)

    # Appendix B
    add_page_break(doc)
    add_kicker(doc, "Appendix B")
    doc.add_heading("Representative event vocabulary", level=1)
    event_rows = [
        ("ContextAssertionAppended", "A validated assertion entered the immutable ledger"),
        ("TopologySnapshotBuilt", "Candidate snapshot and projection manifests are complete"),
        ("TopologySnapshotActivated", "The project active pointer moved atomically"),
        ("ChangeSetResolved", "Base/head differences were semantically normalized"),
        ("ImpactAssessmentCreated", "Canonical explained impact and obligations are immutable"),
        ("WorkOrderAuthorized", "Policy and authority permitted bounded execution"),
        ("ExecutionAttemptCompleted", "A revision-bound attempt produced a signed outcome"),
        ("EvidenceObservationRecorded", "Observed proof or defect was linked to relevant entities"),
        ("PolicyDecisionRecorded", "A deterministic gate or exception was persisted"),
        ("ReleaseAttested", "Approved artifact, environment and evidence were bound"),
        ("OutcomeAttributed", "Incident, escape, rollback or override was linked to prior decisions"),
        ("CandidatePromotionApproved", "A measured configuration/provider/policy version was authorized"),
    ]
    add_table(doc, ["Event", "Meaning"], event_rows, [3200, 6160], first_col_bold=True, font_size=8.7)
    add_callout(doc, "Event rule", "Events carry identifiers and evidence references, not unrestricted source content or credentials. Payload classification and residency are enforced by project policy.", fill=PALE_AMBER, accent=AMBER)

    # Metadata and save
    props = doc.core_properties
    props.title = "Agentic PDLC Context Framework - Target-State Reference Architecture"
    props.subject = "Future-state, vendor-neutral architecture for governed agentic product delivery"
    props.author = "Enterprise Architecture"
    props.keywords = "Agentic PDLC, context fabric, impact intelligence, evidence, plugins, enterprise architecture"
    props.comments = "Target-state leadership review document"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
